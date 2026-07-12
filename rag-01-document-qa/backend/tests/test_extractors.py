import tempfile
import unittest
from pathlib import Path

from docx import Document
import fitz

from services.extraction import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_text,
    clean_text,
)


class ExtractorTests(unittest.TestCase):
    def test_extract_text_from_text_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.txt"
            path.write_text("Hello from text loader", encoding="utf-8")

            self.assertEqual(
                extract_text_from_text(path),
                {
                    "filename": "notes.txt",
                    "pages": [
                        {
                            "page": 1,
                            "text": "Hello from text loader"
                        }
                    ]
                }
            )

    def test_extract_text_from_docx_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.docx"
            document = Document()
            document.add_paragraph("Hello from docx loader")
            document.save(path)

            self.assertEqual(
                extract_text_from_docx(path),
                {
                    "filename": "notes.docx",
                    "pages": [
                        {
                            "page": 1,
                            "text": "Hello from docx loader"
                        }
                    ]
                }
            )

    def test_extract_text_from_pdf_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Hello from pdf loader")
            document.save(path)
            document.close()

            self.assertEqual(
                extract_text_from_pdf(path),
                {
                    "filename": "notes.pdf",
                    "pages": [
                        {
                            "page": 1,
                            "text": "Hello from pdf loader"
                        }
                    ]
                }
            )

    def test_extract_unicode_characters_text_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.txt"
            unicode_text = "Unicode test: 🌟 Café, résumé, こんにちは, 🚀"
            path.write_text(unicode_text, encoding="utf-8")

            self.assertEqual(
                extract_text_from_text(path),
                {
                    "filename": "notes.txt",
                    "pages": [
                        {
                            "page": 1,
                            "text": unicode_text
                        }
                    ]
                }
            )

    def test_extract_unicode_characters_docx_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.docx"
            unicode_text = "Unicode test: Cafe, resume, emoji ✨"
            document = Document()
            document.add_paragraph(unicode_text)
            document.save(path)

            self.assertEqual(
                extract_text_from_docx(path),
                {
                    "filename": "notes.docx",
                    "pages": [
                        {
                            "page": 1,
                            "text": unicode_text
                        }
                    ]
                }
            )

    def test_extract_unicode_characters_pdf_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.pdf"
            unicode_text = "Unicode test: Cafe, resume, check mark"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), unicode_text)
            document.save(path)
            document.close()

            self.assertEqual(
                extract_text_from_pdf(path),
                {
                    "filename": "notes.pdf",
                    "pages": [
                        {
                            "page": 1,
                            "text": unicode_text
                        }
                    ]
                }
            )

    def test_extract_long_pdf_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.pdf"
            document = fitz.open()
            for i in range(1, 51):
                page = document.new_page()
                page.insert_text((72, 72), f"Content of page {i}")
            document.save(path)
            document.close()

            result = extract_text_from_pdf(path)
            self.assertEqual(result["filename"], "notes.pdf")
            self.assertEqual(len(result["pages"]), 50)
            
            for idx, p in enumerate(result["pages"], start=1):
                self.assertEqual(p["page"], idx)
                self.assertEqual(p["text"], f"Content of page {idx}")

    def test_extract_pdf_with_empty_page(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.pdf"
            document = fitz.open()
            page1 = document.new_page()
            page1.insert_text((72, 72), "Page 1 Content")
            page2 = document.new_page()
            page3 = document.new_page()
            page3.insert_text((72, 72), "Page 3 Content")
            document.save(path)
            document.close()

            result = extract_text_from_pdf(path)
            self.assertEqual(
                result,
                {
                    "filename": "notes.pdf",
                    "pages": [
                        {"page": 1, "text": "Page 1 Content"},
                        {"page": 2, "text": ""},
                        {"page": 3, "text": "Page 3 Content"},
                    ]
                }
            )



class TextCleanerTests(unittest.TestCase):
    def test_clean_text_basic(self):
        input_text = "  Hello   World!  "
        self.assertEqual(clean_text(input_text), "Hello World!")

    def test_clean_text_newlines_and_spaces(self):
        input_text = "Line 1\n\n\nLine 2\n\n   \nLine 3"
        expected = "Line 1\n\nLine 2\n\nLine 3"
        self.assertEqual(clean_text(input_text), expected)

    def test_clean_text_carriage_returns(self):
        input_text = "Line 1\r\nLine 2\rLine 3"
        expected = "Line 1\nLine 2\nLine 3"
        self.assertEqual(clean_text(input_text), expected)

    def test_clean_text_empty_input(self):
        self.assertEqual(clean_text(""), "")
        self.assertEqual(clean_text(None), "")


if __name__ == "__main__":
    unittest.main()

